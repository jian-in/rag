"""
文档加载器模块

支持加载多种格式的文档：PDF、TXT、Markdown、DOCX
面试亮点：
  - 使用策略模式根据文件扩展名选择对应的加载器
  - 保留文档元数据（来源文件、页码），方便答案溯源
  - 批量加载支持，方便构建大规模知识库
"""

import os
from pathlib import Path
from typing import List

from langchain_core.documents import Document


class DocumentLoader:
    """
    多格式文档加载器

    支持的格式: .pdf, .txt, .md, .docx
    每个加载的文档都会保留元数据（source, file_name, file_type），
    便于后续在回答中引用来源。
    """

    # 支持的文件格式 → 对应的加载方法名
    SUPPORTED_FORMATS = {".pdf", ".txt", ".md", ".docx"}

    def load_file(self, file_path: str) -> List[Document]:
        """
        加载单个文件，返回 Document 列表

        Args:
            file_path: 文件路径

        Returns:
            List[Document]: 文档列表（PDF 可能包含多个 Document，每页一个）

        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"不支持的文件格式: {suffix}\n"
                f"支持的格式: {', '.join(sorted(self.SUPPORTED_FORMATS))}"
            )

        # 基础元数据
        base_metadata = {
            "source": str(path.resolve()),
            "file_name": path.name,
            "file_type": suffix,
        }

        # 根据文件类型分发到对应的加载方法
        if suffix == ".pdf":
            return self._load_pdf(path, base_metadata)
        elif suffix == ".txt":
            return self._load_text(path, base_metadata)
        elif suffix == ".md":
            return self._load_markdown(path, base_metadata)
        elif suffix == ".docx":
            return self._load_docx(path, base_metadata)

        return []

    def load_directory(self, dir_path: str, recursive: bool = True) -> List[Document]:
        """
        批量加载目录中的所有支持格式的文件

        Args:
            dir_path: 目录路径
            recursive: 是否递归加载子目录

        Returns:
            List[Document]: 所有文档的列表
        """
        dir_path = Path(dir_path)
        if not dir_path.is_dir():
            raise ValueError(f"不是有效的目录: {dir_path}")

        all_docs = []
        pattern = "**/*" if recursive else "*"

        for file_path in sorted(dir_path.glob(pattern)):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                try:
                    docs = self.load_file(str(file_path))
                    all_docs.extend(docs)
                except Exception as e:
                    # 单个文件加载失败不应阻塞整个流程
                    print(f"⚠️  加载失败: {file_path} — {e}")

        print(f"✅ 共加载 {len(all_docs)} 个文档片段，来自 {dir_path}")
        return all_docs

    # ----------------------------------------------------------------
    # 私有方法：各格式的具体加载实现
    # ----------------------------------------------------------------

    def _load_pdf(self, path: Path, base_metadata: dict) -> List[Document]:
        """加载 PDF 文件，每页作为一个 Document"""
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        docs = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                metadata = {**base_metadata, "page": i + 1}
                docs.append(Document(page_content=text.strip(), metadata=metadata))
        return docs

    def _load_text(self, path: Path, base_metadata: dict) -> List[Document]:
        """加载纯文本文件"""
        text = path.read_text(encoding="utf-8")
        if text.strip():
            return [Document(page_content=text.strip(), metadata=base_metadata)]
        return []

    def _load_markdown(self, path: Path, base_metadata: dict) -> List[Document]:
        """加载 Markdown 文件"""
        text = path.read_text(encoding="utf-8")
        if text.strip():
            return [Document(page_content=text.strip(), metadata=base_metadata)]
        return []

    def _load_docx(self, path: Path, base_metadata: dict) -> List[Document]:
        """加载 Word 文档"""
        import docx

        doc = docx.Document(str(path))
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        if text.strip():
            return [Document(page_content=text.strip(), metadata=base_metadata)]
        return []
